from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


PROJECT_NAME = "BillsBuddy"
REPO_URL = "https://github.com/Dudzinska/Wydatki-app.git"
APP_URL_LOCAL = "http://127.0.0.1:8000"
QUESTIONS_PAGE = "/pytania-projektowe"
LOGO_PATH = Path("/workspace/docs/logo_ur.png")
PERSON_1_NAME = "Oliwia Kwasek"
PERSON_2_NAME = "Klaudia Dudzińska"

TECH_ROWS = [
    ("PHP", "8.2.x", "https://www.php.net/"),
    ("Laravel", "12.59.0", "https://laravel.com/"),
    ("Node.js", "22.x", "https://nodejs.org/"),
    ("npm", "10.x", "https://www.npmjs.com/"),
    ("Vite", "7.3.3", "https://vitejs.dev/"),
    ("Tailwind CSS", "3.x", "https://tailwindcss.com/"),
    ("Composer", "2.x", "https://getcomposer.org/"),
    ("SQLite", "3.x", "https://www.sqlite.org/"),
    ("MySQL (opcjonalnie)", "8.0+", "https://www.mysql.com/"),
]


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_title_page(doc: Document, subtitle: str, author_label: str) -> None:
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(5.8))

    doc.add_paragraph()
    title = doc.add_heading("DOKUMENTACJA PROJEKTOWA", level=0)
    title.alignment = 1
    p = doc.add_paragraph(PROJECT_NAME)
    p.alignment = 1
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(20)

    p2 = doc.add_paragraph(subtitle)
    p2.alignment = 1
    p2.runs[0].italic = True

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run(f"Autor: {author_label}\n")
    info.add_run(f"Data: {date.today().isoformat()}\n")
    info.add_run("Kierunek: Informatyka\n")
    info.add_run("Przedmiot: Aplikacje Internetowe")
    info.alignment = 1

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_tech_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Technologia"
    hdr[1].text = "Wersja"
    hdr[2].text = "Oficjalna strona"

    for tech, version, link in TECH_ROWS:
        row = table.add_row().cells
        row[0].text = tech
        row[1].text = version
        row[2].text = link


def add_common_intro(doc: Document) -> None:
    add_heading(doc, "Temat projektu", level=1)
    doc.add_paragraph(
        "BillsBuddy to aplikacja webowa do rozliczania wspólnych wydatków w grupach "
        "(wyjazdy, wspólne mieszkanie, wydarzenia). Użytkownik dodaje grupę, rachunki i pozycje "
        "z paragonu, a system automatycznie oblicza udziały i podpowiada, kto komu i ile oddaje."
    )
    add_bullets(
        doc,
        [
            "Jaki problem rozwiązuje aplikacja? Chaos i błędy w ręcznym rozliczaniu wydatków wielu osób.",
            "Czym się wyróżnia? Automatycznym podziałem kosztów, przejrzystym bilansem oraz planem spłat.",
        ],
    )


def add_developer_runbook(doc: Document) -> None:
    add_heading(doc, "Uruchomienie projektu (developer)", level=1)
    doc.add_paragraph(
        "Poniżej podano dokładne wersje technologii użytych w repozytorium. "
        "Wymagania formalne z dokumentu kursu wskazują PHP 8.5 i Laravel 13; "
        "aktualna implementacja działa na PHP 8.2 i Laravel 12."
    )
    add_tech_table(doc)

    add_heading(doc, "Wymagania programowe", level=2)
    add_bullets(
        doc,
        [
            "System operacyjny: Windows 10/11, Linux lub macOS.",
            "Środowisko uruchomieniowe: PHP 8.2+, Composer 2.x, Node.js 22+, npm 10+.",
            "Silnik bazy danych: SQLite 3 (domyślnie) lub MySQL 8.0+.",
            "Dodatkowe narzędzia: Git, terminal (PowerShell/bash).",
        ],
    )

    add_heading(doc, "Proces instalacji", level=2)
    add_numbered(
        doc,
        [
            f"Pobierz projekt: git clone {REPO_URL}",
            "Przejdź do katalogu: cd Wydatki-app",
            "Zainstaluj backend: composer install",
            "Zainstaluj frontend: npm install",
        ],
    )

    add_heading(doc, "Proces konfiguracji", level=2)
    add_numbered(
        doc,
        [
            "Utwórz plik środowiskowy: skopiuj .env.example do .env",
            "Wygeneruj klucz aplikacji: php artisan key:generate",
            "Skonfiguruj bazę: DB_CONNECTION=sqlite (lub mysql) w .env",
            "Utwórz plik bazy: database/database.sqlite",
            "Wykonaj migracje i seed: php artisan migrate:fresh --seed",
            "Uruchom backend: php artisan serve",
            "Uruchom frontend (drugi terminal): npm run dev",
        ],
    )
    doc.add_paragraph(f"Aplikacja lokalnie: {APP_URL_LOCAL}")

    add_heading(doc, "Dane testowe (seed)", level=2)
    add_bullets(
        doc,
        [
            "admin: oliwia@example.com / password",
            "user: adam@example.com / password",
            "user: ewa@example.com / password",
        ],
    )


def add_user_runbook(doc: Document) -> None:
    add_heading(doc, "Uruchomienie projektu (user)", level=1)
    doc.add_paragraph(
        "Projekt jest przygotowany do uruchamiania lokalnego. "
        "Brak publicznego produkcyjnego deploymentu."
    )
    add_bullets(
        doc,
        [
            "Aplikacja webowa: uruchamiana lokalnie pod adresem http://127.0.0.1:8000.",
            "Wymagania sprzętowe: standardowy laptop/PC, zalecane min. 8 GB RAM.",
            "Wymagana nowoczesna przeglądarka: Chrome, Edge, Firefox.",
        ],
    )


def add_manual_common(doc: Document) -> None:
    add_heading(doc, "Podręcznik użytkownika", level=1)
    add_heading(doc, "Dostęp użytkownika niezalogowanego", level=2)
    add_bullets(
        doc,
        [
            "Użytkownik niezalogowany może wejść do katalogu publicznego grup.",
            "Niezalogowany NIE może tworzyć, edytować ani usuwać zasobów.",
            "Próba operacji modyfikującej przekierowuje do logowania.",
        ],
    )

    add_heading(doc, "Role i uprawnienia (wymaganie 4.0)", level=2)
    add_bullets(
        doc,
        [
            "Administrator: ma panel admina, zarządza profilami użytkowników i ich rolami.",
            "Użytkownik: zarządza własnymi zasobami (grupy, rachunki) w zakresie przyznanego dostępu.",
            "Autoryzacja oparta o middleware i kontrole właściciela zasobu.",
        ],
    )

    add_heading(doc, "Responsywność i zrzuty ekranu", level=2)
    add_bullets(
        doc,
        [
            "Rysunek 1: Dashboard na desktopie.",
            "Rysunek 2: Formularz CRUD na desktopie.",
            "Rysunek 3: Widok mobilny menu i listy danych.",
            "Każdy zrzut powinien mieć podpis: co przedstawia i jaki krok procesu dokumentuje.",
        ],
    )


def add_crud_admin_section(doc: Document) -> None:
    add_heading(doc, "Udokumentowany CRUD zasobu zależnego (wymaganie 3.0/4.0)", level=2)
    doc.add_paragraph(
        "Zasób: Rachunek (Bill) zależny od Grupy (Group). "
        "Relacja: jedna grupa ma wiele rachunków (one-to-many). "
        "Administrator ma pełny dostęp do zasobów w aplikacji."
    )

    add_heading(doc, "CREATE", level=3)
    add_bullets(
        doc,
        [
            "Formularz dodania rachunku zawiera: nazwę wydatku, kwotę, płatnika (select członków grupy).",
            "Walidacja frontend: required, number, min, step, typy pól HTML.",
            "Walidacja backend: required/string/max, numeric/min, weryfikacja płatnika jako członka grupy.",
            "Formularz odwzorowuje relację: płatnik wybierany z listy członków, bez ręcznego wpisywania ID.",
        ],
    )

    add_heading(doc, "READ", level=3)
    add_bullets(
        doc,
        [
            "Lista rachunków ma filtrowanie: tekst, płatnik, zakres kwot.",
            "Sortowanie i paginacja upraszczają pracę na większym zbiorze danych.",
            "Filtry pokrywają realne scenariusze przeglądania historii rozliczeń.",
        ],
    )

    add_heading(doc, "UPDATE", level=3)
    add_bullets(
        doc,
        [
            "Edycja grup i profili ma walidację po stronie klienta i serwera.",
            "Dane po błędzie walidacji wracają do formularza (old input + komunikaty błędów).",
        ],
    )

    add_heading(doc, "DELETE", level=3)
    add_bullets(
        doc,
        [
            "Usuwanie rachunku/grupy realizowane przez dedykowany formularz z metodą DELETE.",
            "Przed usunięciem występuje potwierdzenie akcji w interfejsie.",
            "Autoryzacja serwerowa zabezpiecza usuwanie tylko dla uprawnionych.",
        ],
    )


def add_user_management_40(doc: Document) -> None:
    add_heading(doc, "Rozszerzenie dokumentacji na ocenę 4.0", level=2)
    add_bullets(
        doc,
        [
            "Udokumentowanie ról użytkowników: administrator i użytkownik.",
            "Udokumentowanie uprawnień użytkowników i ograniczeń dostępu.",
            "Udokumentowanie zarządzania zasobami przez użytkowników.",
            "Udokumentowanie zarządzania profilami użytkowników przez administratora.",
        ],
    )


def add_edge_cases(doc: Document) -> None:
    add_heading(doc, "Przypadki brzegowe obsługiwane przez system", level=2)
    add_bullets(
        doc,
        [
            "Nazwa wydatku nie może być samą liczbą.",
            "Kwota wydatku musi być większa od zera.",
            "Liczba sztuk pozycji z paragonu musi być całkowita i >= 1.",
            "Brak dostępu do cudzego zasobu zwraca 403.",
        ],
    )


def add_data_scope(doc: Document) -> None:
    add_heading(doc, "Jakie dane system przechowuje i udostępnia", level=2)
    add_bullets(
        doc,
        [
            "Użytkownicy: imię, e-mail, rola, hasło (hash).",
            "Grupy: nazwa, opis, właściciel, członkowie.",
            "Rachunki: opis, kwota, płatnik, data.",
            "Pozycje rachunków: nazwa, cena, ilość.",
            "Podziały kosztów (splits): udział finansowy każdego członka grupy.",
        ],
    )


def add_main_mechanism(doc: Document) -> None:
    add_heading(doc, "Najważniejszy mechanizm aplikacji", level=2)
    doc.add_paragraph(
        "Po dodaniu rachunku i pozycji z paragonu system automatycznie oblicza udziały "
        "członków grupy, aktualizuje bilans i generuje listę spłat „kto komu i ile oddaje”. "
        "Dodatkowo obsługuje sytuację, gdy suma pozycji paragonu jest różna od kwoty rachunku."
    )


def add_plans(doc: Document) -> None:
    add_heading(doc, "Plany rozbudowy", level=1)
    add_bullets(
        doc,
        [
            "Integracja z płatnościami i oznaczanie spłaconych transferów.",
            "Powiadomienia e-mail/push o nowych rachunkach i zmianach salda.",
            "Eksport raportów do PDF/CSV.",
            "Optymalizacja zapytań i cache dla większej skali danych.",
        ],
    )


def add_questions_reference(doc: Document) -> None:
    add_heading(doc, "Pytania do projektu (3.0 i 4.0)", level=1)
    doc.add_paragraph(
        "Pełny zestaw odpowiedzi znajduje się w aplikacji pod ścieżką "
        f"{QUESTIONS_PAGE} oraz w pliku docs/project_questions_answers.md."
    )


def build_doc_person_1(path: Path) -> None:
    doc = Document()
    set_default_font(doc)
    add_title_page(
        doc,
        subtitle="Dokumentacja wdrożeniowa i administracyjna",
        author_label=f"{PERSON_1_NAME} – część administracyjna i uprawnienia",
    )

    add_heading(doc, f"Autor dokumentacji: {PERSON_1_NAME}", level=2)

    add_common_intro(doc)
    add_developer_runbook(doc)
    add_user_runbook(doc)
    add_manual_common(doc)
    add_crud_admin_section(doc)
    add_user_management_40(doc)
    add_edge_cases(doc)
    add_data_scope(doc)
    add_main_mechanism(doc)
    add_plans(doc)
    add_questions_reference(doc)
    doc.save(path)


def build_doc_person_2(path: Path) -> None:
    doc = Document()
    set_default_font(doc)
    add_title_page(
        doc,
        subtitle="Dokumentacja użytkownika i logiki biznesowej",
        author_label=f"{PERSON_2_NAME} – mechanizmy rozliczeń i user flow",
    )

    add_heading(doc, f"Autor dokumentacji: {PERSON_2_NAME}", level=2)

    add_common_intro(doc)
    add_developer_runbook(doc)
    add_user_runbook(doc)
    add_manual_common(doc)

    add_heading(doc, "Podręcznik użytkownika – user flow (wkład Osoby 2)", level=2)
    add_numbered(
        doc,
        [
            "Zaloguj się i przejdź do „Moje grupy”.",
            "Utwórz grupę i dodaj uczestników (relacja grupa–użytkownicy).",
            "Dodaj rachunek (nazwa, kwota, płatnik).",
            "Dodaj pozycje z paragonu i sprawdź automatyczne przeliczenie sald.",
            "Przejdź do sekcji „Kto ile oddaje” i wykonaj rozliczenie.",
        ],
    )

    add_crud_admin_section(doc)
    add_user_management_40(doc)
    add_edge_cases(doc)
    add_data_scope(doc)
    add_main_mechanism(doc)
    add_plans(doc)
    add_questions_reference(doc)
    doc.save(path)


if __name__ == "__main__":
    build_doc_person_1(Path("/workspace/docs/Dokumentacja_BillsBuddy_Oliwia_Kwasek.docx"))
    build_doc_person_2(Path("/workspace/docs/Dokumentacja_BillsBuddy_Klaudia_Dudzinska.docx"))
    print("Wygenerowano obie dokumentacje DOCX.")
